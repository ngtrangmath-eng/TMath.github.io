from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "SKKN-website-phieu-hoc-tap-Toan-6-2026.docx"


def set_run_font(run, name="Times New Roman", size=13, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_fixed_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    label = paragraph.add_run("Trang ")
    set_run_font(label, size=11)

    begin_run = paragraph.add_run()
    set_run_font(begin_run, size=11)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(fld_begin)

    instr_run = paragraph.add_run()
    set_run_font(instr_run, size=11)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    instr_run._r.append(instr)

    separate_run = paragraph.add_run()
    set_run_font(separate_run, size=11)
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(fld_separate)

    result = paragraph.add_run("1")
    set_run_font(result, size=11)

    end_run = paragraph.add_run()
    set_run_font(end_run, size=11)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    end_run._r.append(fld_end)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.first_line_indent = Inches(0.3)

    for name, size in (("Heading 1", 14), ("Heading 2", 13), ("Heading 3", 13), ("Heading 4", 13)):
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.first_line_indent = None
        style.paragraph_format.space_before = Pt(8 if name == "Heading 1" else 6)
        style.paragraph_format.space_after = Pt(6 if name == "Heading 1" else 4)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(13)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.space_after = Pt(4)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_centered_paragraph(doc, text="", size=13, bold=False, after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_body_paragraph(doc, text, bold_label=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_label:
        r = p.add_run(bold_label)
        set_run_font(r, bold=True)
        r2 = p.add_run(text)
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_section_heading(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    run = p.add_run(text)
    set_run_font(run, size=14, bold=True)
    return p


def add_heading(doc, text, level=2):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = None
    run = p.add_run(text)
    set_run_font(run, size=13 if level >= 2 else 14, bold=True)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.first_line_indent = None
        run = p.add_run(item)
        set_run_font(run)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.first_line_indent = None
        run = p.add_run(item)
        set_run_font(run)


def add_label_block(doc, label, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.3)
    r1 = p.add_run(label)
    set_run_font(r1, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2)
    return p


def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.right_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(2)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_pr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F2F4F7")
        p_pr.append(shd)
        run = p.add_run(line)
        set_run_font(run, name="Courier New", size=10.5)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, size=12, italic=True)


def add_image(doc, image_path, caption, width=6.0):
    path = ROOT / image_path
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_fixed_width(table, widths)
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_shading(hdr[i], "E8EEF5")
        p = hdr[i].paragraphs[0]
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, size=12, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.first_line_indent = None
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            set_run_font(run, size=12)
    set_table_fixed_width(table, widths)
    doc.add_paragraph()
    return table


def add_cover(doc):
    add_centered_paragraph(doc, "ỦY BAN NHÂN DÂN HUYỆN CỦ CHI", size=13, bold=True, after=0)
    add_centered_paragraph(doc, "TRƯỜNG TRUNG HỌC CƠ SỞ THỊ TRẤN 2", size=13, bold=True, after=20)
    add_centered_paragraph(doc, "SÁNG KIẾN KINH NGHIỆM", size=20, bold=True, after=18)
    add_centered_paragraph(
        doc,
        'XÂY DỰNG WEBSITE HỌC TẬP ĐIỆN TỬ MÔN TOÁN 6 "TMATH" NHẰM PHÁT TRIỂN NĂNG LỰC TỰ HỌC CHO HỌC SINH',
        size=16,
        bold=True,
        after=24,
    )
    add_image(doc, "assets/learning-cards-static.png", "Mô hình ba phần học tập trong website phiếu học tập điện tử", width=5.4)
    add_centered_paragraph(doc, "Lĩnh vực: Toán học", size=13, bold=False, after=4)
    add_centered_paragraph(doc, "Người thực hiện: Nguyễn Thị Huyền Trang", size=13, bold=False, after=4)
    add_centered_paragraph(doc, "Năm học: 2026 - 2027", size=13, bold=False, after=90)
    add_centered_paragraph(doc, "Củ Chi, tháng 01 năm 2027", size=13, bold=False, after=0)
    doc.add_page_break()


def add_intro(doc):
    add_section_heading(doc, "I. PHẦN MỞ ĐẦU")
    add_heading(doc, "1. Bối cảnh của đề tài", 2)
    for text in [
        "Trong bối cảnh ngành Giáo dục đang đẩy mạnh đổi mới phương pháp dạy học, đổi mới kiểm tra đánh giá và chuyển đổi số, việc xây dựng học liệu số có khả năng tương tác là yêu cầu thiết thực đối với giáo viên. Chuyển đổi số trong dạy học không chỉ là sử dụng bài giảng trình chiếu, mà còn hướng đến việc tạo ra môi trường học tập giúp học sinh tự học, tự luyện tập, tự kiểm tra và nhận phản hồi kịp thời.",
        "Đối với môn Toán lớp 6, học sinh vừa chuyển từ bậc Tiểu học lên Trung học cơ sở nên còn nhiều bỡ ngỡ về cách học, cách ghi nhớ kiến thức, cách trình bày bài làm và tự đánh giá mức độ hiểu bài. Các nội dung nền tảng như tập hợp, số tự nhiên, phép tính với số tự nhiên, lũy thừa và thứ tự thực hiện phép tính đòi hỏi học sinh phải được ôn luyện thường xuyên, có ví dụ rõ ràng và có cơ hội kiểm tra lỗi sai ngay trong quá trình học.",
        "Thực tế giảng dạy cho thấy một số học sinh còn thiếu hứng thú học Toán, chưa có thói quen tự học sau giờ lên lớp, dễ nhầm lẫn khái niệm, kí hiệu và quy trình giải toán. Phiếu bài tập giấy vẫn cần thiết nhưng còn hạn chế ở khả năng phản hồi tức thời, theo dõi tiến độ và hỗ trợ học sinh học theo tốc độ riêng.",
        "Từ thực tiễn đó, tôi xây dựng website phiếu học tập điện tử môn Toán 6 với sự hỗ trợ của trí tuệ nhân tạo. Website được thiết kế như một môi trường học tập trực tuyến đơn giản, dễ truy cập, gồm trang chủ, danh sách bài học, hướng dẫn học tập, tài khoản học sinh, theo dõi tiến độ và các phiếu học tập tương tác cho từng bài. Giáo viên sử dụng prompt mẫu để tạo hoặc cập nhật phiếu học tập từ các file nội dung đã chuẩn bị, qua đó tiết kiệm thời gian và bảo đảm cấu trúc học liệu thống nhất.",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "2. Lý do chọn đề tài", 2)
    for text in [
        "Tôi lựa chọn đề tài này trước hết xuất phát từ yêu cầu đổi mới phương pháp dạy học môn Toán theo định hướng phát triển phẩm chất, năng lực học sinh. Trong dạy học hiện nay, học sinh không chỉ cần ghi nhớ kiến thức mà còn phải biết tự học, biết luyện tập, biết phát hiện lỗi sai và biết vận dụng kiến thức vào tình huống cụ thể. Vì vậy, giáo viên cần có thêm công cụ hỗ trợ để tổ chức hoạt động học tập theo hướng chủ động, thường xuyên và có phản hồi.",
        "Môn Toán lớp 6 giữ vai trò nền tảng trong toàn bộ quá trình học Toán ở bậc Trung học cơ sở. Đây là giai đoạn học sinh chuyển từ Tiểu học lên THCS, bắt đầu làm quen với cách học theo từng môn, từng giáo viên bộ môn và yêu cầu tự học cao hơn. Nếu học sinh chưa nắm vững các kiến thức đầu cấp như tập hợp, số tự nhiên, phép tính, lũy thừa, thứ tự thực hiện phép tính thì các em sẽ gặp khó khăn khi học những nội dung tiếp theo như số nguyên, phân số, số thập phân và đại số sau này.",
        "Qua thực tế giảng dạy, tôi nhận thấy một bộ phận học sinh lớp 6 còn học Toán theo thói quen thụ động. Nhiều em chỉ nghe giảng trên lớp, ghi bài và làm bài khi giáo viên yêu cầu, nhưng chưa biết tự xem lại kiến thức trọng tâm, chưa biết tự luyện tập thêm và chưa có thói quen kiểm tra xem mình đã hiểu bài đến đâu. Khi gặp bài sai, một số em chỉ sửa đáp án mà chưa biết nguyên nhân sai ở khái niệm, quy tắc hay bước trình bày.",
        "Bên cạnh đó, học sinh lớp 6 rất cần các hình thức học tập trực quan, ngắn gọn và có tính tương tác. Nếu chỉ sử dụng sách giáo khoa, vở ghi hoặc phiếu bài tập giấy, học sinh thường phải chờ giáo viên chấm, chưa nhận được phản hồi ngay sau khi làm bài. Đối với những em học chậm, việc không có gợi ý kịp thời dễ làm các em nản; đối với học sinh khá giỏi, phiếu giấy thông thường chưa tạo được nhiều cơ hội tự luyện tập, tự kiểm tra và mở rộng theo tốc độ riêng.",
        "Từ phía giáo viên, việc chuẩn bị phiếu học tập chất lượng cho từng bài học cũng là một yêu cầu mất nhiều thời gian. Một phiếu học tập hiệu quả cần có mục tiêu, kiến thức trọng tâm, ví dụ mẫu, bài luyện tập, câu hỏi kiểm tra, đáp án, thang điểm và nhận xét. Nếu giáo viên tự thiết kế từng phiếu bằng cách thủ công, mỗi bài có một kiểu trình bày khác nhau, quá trình biên soạn dễ mất nhiều công sức, khó cập nhật, khó chia sẻ và khó nhân rộng trong tổ chuyên môn.",
        "Trong bối cảnh chuyển đổi số, việc xây dựng website phiếu học tập điện tử là hướng đi phù hợp với điều kiện thực tế. Website giúp gom các bài học vào một địa chỉ học tập thống nhất, học sinh có thể truy cập khi học trên lớp, ôn tập ở nhà hoặc chuẩn bị trước bài mới. Các chức năng như danh sách bài học, tìm kiếm, hướng dẫn học tập, bài đã mở, bài gần nhất, chuỗi ngày học và huy hiệu nỗ lực giúp việc học Toán trở nên rõ lộ trình hơn, có động lực hơn và dễ theo dõi hơn.",
        "Sản phẩm website phiếu học tập điện tử Toán 6 mà tôi xây dựng có cấu trúc phù hợp với quá trình học của học sinh: Phần A - Kiến thức trọng tâm, Phần B - Luyện tập, Phần C - Kiểm tra đánh giá. Cấu trúc này giúp học sinh không học rời rạc mà đi theo quy trình: nắm kiến thức, làm bài, nhận phản hồi, tự kiểm tra và điều chỉnh lỗi sai. Đây là điểm cần thiết để hỗ trợ học sinh lớp 6 hình thành năng lực tự học từng bước, phù hợp với đặc điểm tâm lý lứa tuổi đầu cấp.",
        "Một lý do quan trọng khác là nhu cầu hỗ trợ giáo viên tạo học liệu bằng trí tuệ nhân tạo nhưng vẫn bảo đảm kiểm soát chuyên môn. Nếu chỉ dùng AI bằng những yêu cầu chung chung, sản phẩm học liệu có thể thiếu thống nhất, sai lệch nội dung hoặc chưa phù hợp với học sinh. Vì vậy, tôi xây dựng prompt mẫu và quy trình file nguồn để giáo viên chủ động chuẩn bị nội dung chuyên môn, sau đó dùng AI hỗ trợ chuyển hóa thành phiếu HTML tương tác. Cách làm này giữ vai trò quyết định của giáo viên ở khâu lựa chọn kiến thức, câu hỏi, đáp án và thang điểm.",
        "Việc lựa chọn đề tài “Xây dựng website phiếu học tập điện tử môn Toán 6 nhằm phát triển năng lực tự học cho học sinh” vì thế có ý nghĩa thiết thực. Đề tài vừa giải quyết nhu cầu học tập của học sinh, vừa hỗ trợ công việc chuẩn bị bài của giáo viên, đồng thời góp phần thực hiện chuyển đổi số trong nhà trường bằng một sản phẩm cụ thể, dễ sử dụng, dễ cập nhật và có khả năng mở rộng cho nhiều bài học khác.",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "3. Phạm vi và đối tượng nghiên cứu", 2)
    for text in [
        "Đề tài tập trung nghiên cứu việc thiết kế, xây dựng và sử dụng website phiếu học tập điện tử trong dạy học môn Toán lớp 6. Sản phẩm trọng tâm là website phiếu học tập Toán 6 Chương 1, gồm mục lục chương, các bài học từ Bài 1 đến Bài 7 và phiếu ôn tập chương.",
        "Đối tượng áp dụng là học sinh lớp 6 tại Trường THCS Thị Trấn 2, đặc biệt là những học sinh cần được hỗ trợ trong quá trình tự học, ôn tập, luyện tập và kiểm tra kiến thức. Đề tài đồng thời hướng đến giáo viên môn Toán có nhu cầu xây dựng học liệu số nhanh, thống nhất, dễ cập nhật và có thể chia sẻ cho học sinh.",
        "Nội dung triển khai trước hết tập trung vào Chương 1 - Số học, gồm các bài: Tập hợp; Cách ghi số tự nhiên; Thứ tự trong tập hợp các số tự nhiên; Phép cộng và phép trừ số tự nhiên; Phép nhân và phép chia số tự nhiên; Lũy thừa với số mũ tự nhiên; Thứ tự thực hiện các phép tính; Ôn tập chương 1.",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "4. Mục đích nghiên cứu", 2)
    add_body_paragraph(
        doc,
        "Đề tài nhằm xây dựng một website phiếu học tập điện tử môn Toán 6 để hỗ trợ học sinh tự học, tự luyện tập, tự kiểm tra kiến thức và hình thành thói quen học tập chủ động."
    )
    add_bullets(
        doc,
        [
            "Hỗ trợ học sinh tiếp cận kiến thức theo lộ trình rõ ràng: kiến thức trọng tâm, luyện tập, kiểm tra đánh giá.",
            "Tăng cơ hội phản hồi cho học sinh thông qua câu hỏi tương tác, gợi ý, đáp án và nhận xét kết quả.",
            "Hỗ trợ giáo viên thiết kế phiếu học tập bằng prompt mẫu, giảm thời gian định dạng, chuẩn hóa cấu trúc học liệu và dễ nhân rộng cho nhiều bài học.",
            "Góp phần đổi mới phương pháp dạy học, tăng cường ứng dụng công nghệ thông tin và thực hiện chuyển đổi số trong nhà trường.",
        ],
    )

    add_heading(doc, "5. Điểm mới trong kết quả nghiên cứu", 2)
    for text in [
        "Điểm mới của đề tài là sản phẩm không chỉ là một file bài giảng hoặc một phiếu bài tập đơn lẻ, mà là một website học tập có cấu trúc hoàn chỉnh. Trang chủ gom các phiếu học tập Toán 6, có video giới thiệu, danh sách bài học, tìm kiếm, lọc bài, hướng dẫn học tập, đăng nhập học sinh và theo dõi tiến độ.",
        "Mỗi phiếu học tập được thiết kế theo ba phần thống nhất: Phần A - Kiến thức trọng tâm, Phần B - Luyện tập, Phần C - Kiểm tra đánh giá. Cấu trúc này giúp học sinh biết rõ cần học gì, luyện tập như thế nào và kiểm tra mức độ hiểu bài ra sao.",
        "Điểm mới quan trọng khác là xây dựng quy trình tạo phiếu dựa trên prompt mẫu. Giáo viên chuẩn bị nội dung chuyên môn trong các file Markdown, sau đó dùng lệnh mẫu để AI hỗ trợ chuyển thành HTML tương tác. Nhờ vậy, sản phẩm vừa giữ được vai trò kiểm duyệt chuyên môn của giáo viên, vừa tận dụng được khả năng hỗ trợ của công nghệ.",
        "Website còn có khả năng lưu tiến độ học tập trên thiết bị, ghi nhận bài đã mở, bài học gần nhất, chuỗi ngày học và huy hiệu nỗ lực. Những yếu tố này giúp học sinh có thêm động lực tự học, đồng thời hỗ trợ giáo viên quan sát quá trình học tập khi tổ chức nhiệm vụ trên lớp hoặc ở nhà.",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "6. Thời gian nghiên cứu", 2)
    add_body_paragraph(doc, "Từ ngày 01 tháng 7 năm 2026 đến ngày 15 tháng 01 năm 2027.")


def add_theory_and_status(doc):
    add_section_heading(doc, "II. NỘI DUNG")
    add_heading(doc, "1. Cơ sở lí luận của đề tài", 2)
    theory_sections = [
        (
            "1.1. Đổi mới giáo dục phổ thông và chuyển đổi số",
            [
                "Chương trình Giáo dục phổ thông 2018 định hướng phát triển phẩm chất, năng lực của học sinh, trong đó học sinh cần biết tự học, vận dụng kiến thức và chủ động giải quyết vấn đề. Vì vậy, giáo viên cần tổ chức dạy học theo hướng học sinh được tham gia, luyện tập, phản hồi và tự đánh giá.",
                "Chuyển đổi số trong giáo dục tạo điều kiện để giáo viên xây dựng học liệu số, tổ chức hoạt động học tập linh hoạt và mở rộng không gian học tập ra ngoài lớp học. Website phiếu học tập điện tử là một hình thức học liệu số phù hợp vì có thể kết hợp văn bản, hình ảnh, bài tập tương tác, kiểm tra và theo dõi tiến độ.",
            ],
        ),
        (
            "1.2. Năng lực tự học của học sinh lớp 6",
            [
                "Tự học là năng lực quan trọng giúp học sinh xác định mục tiêu học tập, lựa chọn nội dung cần học, thực hiện nhiệm vụ và tự đánh giá kết quả. Với học sinh lớp 6, năng lực này cần được hình thành từng bước vì các em đang chuyển sang môi trường học tập mới ở bậc Trung học cơ sở.",
                "Một công cụ học tập tốt cần hướng dẫn học sinh theo trình tự dễ hiểu: đọc mục tiêu, nắm kiến thức trọng tâm, xem ví dụ, luyện tập, kiểm tra và điều chỉnh lỗi sai. Website phiếu học tập điện tử đáp ứng yêu cầu này nhờ cấu trúc bài học rõ ràng và khả năng phản hồi ngay trong quá trình học.",
            ],
        ),
        (
            "1.3. Đặc điểm môn Toán lớp 6",
            [
                "Toán 6 có nhiều kiến thức nền tảng, liên quan chặt chẽ với nhau. Nếu học sinh chưa hiểu vững khái niệm ban đầu, các em dễ gặp khó khăn ở những nội dung tiếp theo. Do đó, việc hệ thống hóa kiến thức và tạo cơ hội luyện tập thường xuyên là rất cần thiết.",
                "Các bài học về tập hợp, số tự nhiên, phép tính, lũy thừa và thứ tự thực hiện phép tính có thể được chuyển hóa thành phiếu học tập điện tử với ví dụ mẫu, câu hỏi trắc nghiệm, bài tập nhập đáp án, kiểm tra nhanh và nhận xét kết quả. Cách làm này giúp kiến thức Toán trở nên trực quan hơn, phù hợp với đặc điểm nhận thức của học sinh đầu cấp.",
            ],
        ),
        (
            "1.4. Vai trò của AI và prompt mẫu trong hỗ trợ giáo viên",
            [
                "AI có thể hỗ trợ giáo viên trong việc chuyển đổi nội dung bài học thành phiếu học tập có cấu trúc, gợi ý bài tập, tạo phản hồi, sắp xếp giao diện và xuất bản thành HTML. Tuy nhiên, nội dung chuyên môn, đáp án, thang điểm và mức độ phù hợp vẫn cần giáo viên kiểm tra, điều chỉnh.",
                "Prompt mẫu giúp quá trình sử dụng AI trở nên có kiểm soát. Thay vì yêu cầu chung chung, giáo viên sử dụng lệnh thống nhất, cung cấp đúng nguồn dữ liệu và yêu cầu đầu ra rõ ràng. Nhờ đó, các phiếu học tập có cùng cấu trúc, dễ kiểm tra, dễ cập nhật và hạn chế sai lệch nội dung.",
            ],
        ),
    ]
    for heading, paragraphs in theory_sections:
        add_heading(doc, heading, 3)
        for text in paragraphs:
            add_body_paragraph(doc, text)

    add_heading(doc, "2. Thực trạng vấn đề", 2)
    add_heading(doc, "2.1. Thuận lợi", 3)
    add_bullets(
        doc,
        [
            "Nhà trường quan tâm đến đổi mới phương pháp dạy học và khuyến khích giáo viên ứng dụng công nghệ thông tin trong giảng dạy.",
            "Giáo viên có kinh nghiệm giảng dạy Toán 6, nắm được những nội dung học sinh thường nhầm lẫn và có thể lựa chọn kiến thức phù hợp để đưa vào phiếu học tập.",
            "Nhiều học sinh đã quen với thiết bị số, có thể mở đường link, xem video, làm bài trực tuyến và học lại bài ở nhà dưới sự hướng dẫn của giáo viên, phụ huynh.",
            "Các công cụ hỗ trợ tạo học liệu, lưu trữ website tĩnh và AI ngày càng phổ biến, tạo điều kiện để giáo viên xây dựng sản phẩm học tập linh hoạt, chi phí thấp.",
        ],
    )
    add_heading(doc, "2.2. Khó khăn", 3)
    add_bullets(
        doc,
        [
            "Một số học sinh chưa có thói quen tự học, còn học thụ động và dễ bỏ qua bước tự kiểm tra sau khi học bài.",
            "Điều kiện thiết bị, đường truyền internet và kỹ năng sử dụng công nghệ của học sinh chưa đồng đều.",
            "Giáo viên mất nhiều thời gian nếu phải thiết kế từng phiếu học tập thủ công, nhất là khi cần có giao diện tương tác, phản hồi và kiểm tra đánh giá.",
            "Nếu không có quy trình kiểm soát, việc dùng AI để tạo học liệu có thể dẫn đến nội dung thiếu chính xác, sai thang điểm hoặc chưa phù hợp với đặc điểm học sinh lớp 6.",
        ],
    )
    add_body_paragraph(
        doc,
        "Từ thực trạng trên, cần có một giải pháp vừa hỗ trợ học sinh tự học, vừa hỗ trợ giáo viên tạo học liệu số nhanh nhưng vẫn bảo đảm kiểm soát chuyên môn. Website phiếu học tập điện tử kết hợp prompt mẫu là giải pháp phù hợp với điều kiện thực tế.",
    )


def add_measures(doc):
    add_heading(doc, "3. Các biện pháp", 2)
    add_body_paragraph(
        doc,
        "Trên cơ sở nghiên cứu thực trạng và đặc điểm học sinh lớp 6, tôi xây dựng website phiếu học tập điện tử môn Toán 6 theo hướng lấy học sinh làm trung tâm, đồng thời thiết kế prompt mẫu để hỗ trợ giáo viên tạo học liệu. Các biện pháp được triển khai theo nguyên tắc: rõ mục đích, rõ cách thực hiện, có sản phẩm cụ thể và có khả năng điều chỉnh sau khi sử dụng.",
    )

    rows = [
        (
            "Xây dựng website học tập",
            "Tạo trang chủ, danh sách bài học, điều hướng, hướng dẫn học tập, video giới thiệu và hệ thống phiếu HTML độc lập.",
            "Học sinh dễ truy cập; giáo viên dễ chia sẻ, cập nhật và nhân rộng.",
        ),
        (
            "Chuẩn hóa phiếu ba phần",
            "Mỗi bài gồm Phần A - kiến thức trọng tâm, Phần B - luyện tập, Phần C - kiểm tra đánh giá.",
            "Học sinh học theo lộ trình; giáo viên kiểm soát đủ kiến thức, luyện tập và đánh giá.",
        ),
        (
            "Dùng prompt mẫu",
            "Giáo viên chuẩn bị lesson.md và ba file A/B/C, sau đó dùng lệnh mẫu để tạo hoặc cập nhật HTML.",
            "Giảm thời gian thiết kế, thống nhất cấu trúc, hạn chế sai lệch khi dùng AI.",
        ),
        (
            "Theo dõi tiến độ",
            "Website lưu bài đã mở, bài gần nhất, chuỗi ngày học, huy hiệu nỗ lực và khóa phần kiểm tra khi chưa đăng nhập.",
            "Tăng động lực tự học; hỗ trợ giáo viên tổ chức giao nhiệm vụ và kiểm tra.",
        ),
    ]
    add_table(doc, ["Biện pháp", "Cách thực hiện chính", "Tác dụng"], rows, [2100, 4300, 2700])

    add_heading(doc, "3.1. Biện pháp 1: Xây dựng website phiếu học tập điện tử Toán 6", 3)
    add_label_block(doc, "Mục đích: ", "Tạo một môi trường học tập trực tuyến đơn giản, trực quan, dễ sử dụng, giúp học sinh có thể học bài, luyện tập và kiểm tra kiến thức ở bất cứ thời điểm phù hợp.")
    add_label_block(doc, "Môi trường làm việc sử dụng: ", "Trong quá trình xây dựng sản phẩm, tôi sử dụng kết hợp ChatGPT và Codex. ChatGPT hỗ trợ trao đổi bằng ngôn ngữ tự nhiên, hình thành ý tưởng sư phạm và xây dựng prompt; Codex hỗ trợ làm việc trực tiếp với mã nguồn website, chỉnh sửa file, kiểm tra sản phẩm và tạo tài liệu minh họa.")
    add_table(
        doc,
        ["Môi trường", "Cách sử dụng trong sáng kiến", "Kết quả đạt được"],
        [
            (
                "ChatGPT",
                "Trao đổi ý tưởng, xác định mục tiêu bài học, xây dựng cấu trúc phiếu học tập, viết và điều chỉnh prompt mẫu, rà soát cách diễn đạt nội dung sư phạm.",
                "Có khung ý tưởng, prompt mẫu và định hướng nội dung phù hợp với học sinh lớp 6.",
            ),
            (
                "Codex",
                "Làm việc trực tiếp trong thư mục dự án website; đọc, tạo và chỉnh sửa các file HTML, CSS, JavaScript, Markdown; kiểm tra cấu trúc bài học; chụp ảnh minh họa từ website; tạo và cập nhật file SKKN.",
                "Hoàn thiện trang chủ, danh sách bài học, phiếu HTML tương tác và tài liệu SKKN có hình minh họa sản phẩm.",
            ),
            (
                "Workspace dự án",
                "Lưu toàn bộ sản phẩm trong thư mục cTrang_sangkien, gồm index.html, assets, templates và các thư mục Toan6-Chuong1-BaiX.",
                "Dữ liệu bài học được quản lý rõ ràng, dễ cập nhật và dễ nhân rộng.",
            ),
            (
                "Trình duyệt",
                "Mở trực tiếp file HTML để kiểm tra giao diện, điều hướng, phản hồi bài tập, nút nộp bài, kết quả và nút in.",
                "Bảo đảm website có thể sử dụng thực tế trên máy tính và điện thoại.",
            ),
        ],
        [1900, 4700, 2500],
    )
    add_label_block(doc, "Cách thực hiện: ", "Tôi trình bày quy trình thực hiện theo các bước cụ thể dưới đây. Mỗi bước đều có nội dung công việc và câu lệnh/prompt đã sử dụng để tạo sản phẩm website phiếu học tập điện tử.")
    add_label_block(doc, "Bước 1. Xác định mục tiêu và cấu trúc website: ", "Tôi xác định website cần phục vụ trực tiếp cho học sinh lớp 6 tự học môn Toán, vì vậy website phải có trang chủ, danh sách bài học, hướng dẫn học tập, đăng nhập học sinh, theo dõi tiến độ và các phiếu học tập tương tác.")
    add_code_block(
        doc,
        [
            "Câu lệnh/prompt đã thực hiện:",
            "Xây dựng website phiếu học tập điện tử môn Toán 6, có trang chủ, danh sách bài học, hướng dẫn học tập, đăng nhập học sinh, theo dõi tiến độ và các phiếu học tập tương tác.",
        ],
    )
    add_label_block(doc, "Bước 2. Thiết kế trang chủ website: ", "Tôi xây dựng file index.html làm trang chủ. Trang chủ có tên website, lời chào học sinh, video giới thiệu, ba hoạt động chính, danh sách bài học, thanh tìm kiếm, bộ lọc chương, khu vực đăng nhập và phần thành tích học tập.")
    add_code_block(
        doc,
        [
            "Câu lệnh/prompt đã thực hiện:",
            "Tạo trang chủ Phiếu học tập Toán 6 với thanh điều hướng gồm Trang chủ, Vào học, Thành tích, Đăng nhập, Hướng dẫn; có video giới thiệu, danh sách bài học Chương 1 và khu vực theo dõi tiến độ học tập.",
        ],
    )
    add_label_block(doc, "Bước 3. Tạo các file dùng chung cho giao diện và chức năng: ", "Tôi tạo thư mục assets để lưu các file dùng chung như site.css, site.js, lesson-shell.css và lesson-shell.js. Các file này giúp thống nhất giao diện, xử lý tìm kiếm, lọc bài, đăng nhập, lưu tiến độ, ghi nhớ bài học gần nhất và tạo thanh điều hướng trong từng bài.")
    add_code_block(
        doc,
        [
            "Câu lệnh/prompt đã thực hiện:",
            "Tạo các file dùng chung cho website: assets/site.css, assets/site.js, assets/lesson-shell.css, assets/lesson-shell.js để thống nhất giao diện, xử lý tìm kiếm, đăng nhập, lưu tiến độ và điều hướng bài học.",
        ],
    )
    add_label_block(doc, "Bước 4. Chuẩn bị cấu trúc thư mục và file nguồn cho từng bài: ", "Mỗi bài học được đặt trong một thư mục riêng. Trong thư mục có lesson.md để ghi thông tin bài học, file PhanA cho kiến thức trọng tâm, file PhanB cho luyện tập, file PhanC cho kiểm tra đánh giá và file HTML đầu ra để học sinh sử dụng.")
    add_code_block(
        doc,
        [
            "Câu lệnh/prompt đã thực hiện:",
            "Tạo cấu trúc bài học Toan6-Chuong1-BaiX gồm lesson.md, BaiX-PhanA.md, BaiX-PhanB.md, BaiX-PhanC.md và file HTML đầu ra.",
        ],
    )
    add_label_block(doc, "Bước 5. Biên soạn nội dung chuyên môn cho từng phần: ", "Giáo viên chuẩn bị nội dung trước khi yêu cầu AI tạo phiếu. Phần A gồm mục tiêu, ghi nhớ, ví dụ mẫu, lỗi thường gặp; Phần B gồm bài luyện tập, gợi ý, phản hồi; Phần C gồm câu hỏi kiểm tra, đáp án, thang điểm và nhận xét kết quả.")
    add_code_block(
        doc,
        [
            "Câu lệnh/prompt đã thực hiện:",
            "Biên soạn Phần A - Kiến thức trọng tâm, Phần B - Luyện tập, Phần C - Kiểm tra đánh giá cho bài học; giữ đúng kiến thức, đáp án và thang điểm.",
        ],
    )
    add_label_block(doc, "Bước 6. Dùng câu lệnh mẫu để tạo phiếu học tập điện tử: ", "Sau khi chuẩn bị đủ file nguồn, tôi dùng câu lệnh mẫu để AI đọc đúng thư mục bài học và tạo file HTML tương tác. Đây là bước chuyển nội dung chuyên môn thành sản phẩm học tập điện tử.")
    add_code_block(
        doc,
        [
            "Các câu lệnh đã thực hiện:",
            "Hoàn thành phiếu học tập Toan6-Chuong1-Bai1",
            "Hoàn thành phiếu học tập Toan6-Chuong1-Bai2",
            "Hoàn thành phiếu học tập Toan6-Chuong1-Bai3",
            "Hoàn thành phiếu học tập Toan6-Chuong1-Bai4",
            "Hoàn thành phiếu học tập Toan6-Chuong1-Bai5",
            "Hoàn thành phiếu học tập Toan6-Chuong1-Bai6",
            "Hoàn thành phiếu học tập Toan6-Chuong1-Bai7",
            "Hoàn thành phiếu học tập Toan6-Chuong1-OnTap",
        ],
    )
    add_label_block(doc, "Bước 7. Kiểm tra, chỉnh sửa và hoàn thiện website: ", "Tôi mở từng file HTML bằng trình duyệt để kiểm tra đủ ba phần A/B/C, thử câu hỏi luyện tập, thử phản hồi đúng sai, kiểm tra nút nộp bài, kết quả, nút in và thanh điều hướng. Nếu nội dung chưa rõ hoặc giao diện chưa phù hợp, tôi điều chỉnh lại prompt hoặc file nguồn rồi tạo lại phiếu.")
    add_code_block(
        doc,
        [
            "Câu lệnh/prompt đã thực hiện:",
            "Kiểm tra lại phiếu học tập HTML: đủ Phần A, Phần B, Phần C; đúng kiến thức; đúng đáp án; đúng thang điểm; có phản hồi; có nút in; giao diện dễ đọc trên máy tính và điện thoại.",
        ],
    )
    add_label_block(doc, "Bước 8. Hướng dẫn học sinh sử dụng website: ", "Sau khi hoàn thiện, giáo viên chia sẻ website cho học sinh bằng đường link hoặc mã QR. Học sinh học theo trình tự: đăng nhập, chọn bài học, đọc Phần A, làm Phần B, hoàn thành Phần C, xem kết quả và ghi lại nội dung chưa hiểu để hỏi giáo viên.")
    add_code_block(
        doc,
        [
            "Câu lệnh/prompt hướng dẫn học sinh:",
            "Em hãy vào website Phiếu học tập Toán 6, chọn bài học được giao, học lần lượt Phần A - Kiến thức trọng tâm, Phần B - Luyện tập, Phần C - Kiểm tra đánh giá, sau đó xem kết quả và ghi lại câu chưa hiểu.",
        ],
    )
    add_bullets(
        doc,
        [
            "Trang chủ giới thiệu mục tiêu học tập, video hướng dẫn, ba hoạt động chính và danh sách bài học.",
            "Danh sách bài học gồm mục lục chương 1, các bài từ Bài 1 đến Bài 7 và phiếu ôn tập chương.",
            "Thanh điều hướng chung trong từng bài cho phép học sinh quay về trang chủ, chuyển bài trước, bài sau và in phiếu khi cần.",
            "Website hoạt động dưới dạng trang tĩnh nên dễ chia sẻ qua đường link, dễ triển khai trên GitHub Pages hoặc mở trực tiếp trên máy tính.",
        ],
    )
    add_label_block(doc, "Hình ảnh minh họa cách thực hiện: ", "Các hình dưới đây được chụp trực tiếp từ sản phẩm website học tập điện tử môn Toán 6 TMath đã xây dựng.")
    add_image(doc, "skkn_assets/website-trang-chu.png", "Minh họa Biện pháp 1. Trang chủ website học tập điện tử TMath môn Toán 6", width=5.8)
    add_image(doc, "skkn_assets/website-danh-sach-bai.png", "Minh họa Biện pháp 1. Danh sách bài học Chương 1 trên website TMath", width=5.8)
    add_image(doc, "skkn_assets/website-phieu-bai-1.png", "Minh họa Biện pháp 1. Phiếu học tập điện tử Bài 1 - Tập hợp trên website TMath", width=5.8)
    add_label_block(doc, "Sản phẩm đạt được: ", "Website học tập điện tử TMath môn Toán 6 Chương 1 có giao diện thống nhất, dễ truy cập, hỗ trợ học sinh học theo từng bài và tạo tiền đề để mở rộng sang các chương khác.")

    add_heading(doc, "3.2. Biện pháp 2: Chuẩn hóa cấu trúc mỗi phiếu học tập theo ba phần", 3)
    add_label_block(doc, "Mục đích: ", "Giúp học sinh không học rời rạc mà đi theo một quy trình ổn định: nắm kiến thức, luyện tập, kiểm tra và tự điều chỉnh.")
    add_label_block(doc, "Môi trường làm việc sử dụng trong biện pháp 2: ", "Biện pháp này được thực hiện trong môi trường ChatGPT, Codex và trình duyệt. ChatGPT hỗ trợ thiết kế cấu trúc nội dung, viết prompt cho từng phần của phiếu học tập; Codex hỗ trợ làm việc trực tiếp với các file HTML, CSS, JavaScript của website; trình duyệt dùng để kiểm tra sản phẩm sau khi tạo.")
    add_table(
        doc,
        ["Môi trường", "Cách sử dụng trong biện pháp 2", "Sản phẩm thu được"],
        [
            (
                "ChatGPT",
                "Trao đổi để xác định cấu trúc chuẩn của một phiếu học tập; viết prompt cho từng phần A, B, C; rà soát cách diễn đạt mục tiêu, ghi nhớ, câu hỏi, gợi ý, đáp án và nhận xét.",
                "Có khung phiếu học tập thống nhất, phù hợp với học sinh lớp 6 và định hướng phát triển năng lực tự học.",
            ),
            (
                "Codex",
                "Đọc các file nội dung nguồn, chỉnh sửa file HTML/CSS/JavaScript, tạo các tab Phần A - Phần B - Phần C - Kết quả, kiểm tra nút phản hồi, nút nộp bài và nút in.",
                "Các phiếu HTML có cùng cấu trúc, cùng cách điều hướng và cùng logic phản hồi.",
            ),
            (
                "Trình duyệt",
                "Mở phiếu học tập để kiểm tra trực tiếp từng tab, thử trả lời câu hỏi, xem điểm, xem nhận xét và kiểm tra giao diện trên màn hình máy tính.",
                "Bảo đảm học sinh dễ thao tác, giáo viên dễ hướng dẫn và dễ kiểm tra sản phẩm.",
            ),
        ],
        [1900, 4700, 2500],
    )
    add_label_block(doc, "Cách thực hiện: ", "Tôi chuẩn hóa mỗi phiếu học tập theo các bước tuần tự dưới đây. Mỗi bước đều nêu rõ việc đã làm trong môi trường ChatGPT, Codex hoặc trình duyệt, kèm câu lệnh/prompt đã thực hiện để giáo viên có thể lặp lại khi tạo phiếu cho bài học mới.")
    add_label_block(doc, "Bước 1. Xác định cấu trúc bắt buộc của một phiếu học tập điện tử: ", "Trước khi tạo nội dung chi tiết, tôi quy định mọi phiếu học tập đều phải có cùng bố cục: thông tin bài học, Phần A - Kiến thức trọng tâm, Phần B - Luyện tập, Phần C - Kiểm tra đánh giá và phần Kết quả. Cấu trúc cố định giúp học sinh không bị lúng túng khi chuyển từ bài này sang bài khác.")
    add_code_block(
        doc,
        [
            "Câu lệnh/prompt đã thực hiện:",
            "Chuẩn hóa mỗi phiếu học tập Toán 6 theo cấu trúc: thông tin bài học, Phần A - Kiến thức trọng tâm, Phần B - Luyện tập, Phần C - Kiểm tra đánh giá, phần Kết quả; bảo đảm học sinh có thể tự học theo đúng trình tự.",
        ],
    )
    add_label_block(doc, "Bước 2. Thiết kế Phần A - Kiến thức trọng tâm: ", "Phần A được dùng để học sinh tự đọc trước khi làm bài. Nội dung cần ngắn gọn, chính xác, có mục tiêu bài học, kiến thức cần nhớ, ví dụ mẫu, lỗi thường gặp và câu hỏi tự kiểm tra nhanh. Cách trình bày này giúp học sinh biết mình cần học gì và cần tránh những sai lầm nào.")
    add_code_block(
        doc,
        [
            "Câu lệnh/prompt đã thực hiện:",
            "Viết Phần A - Kiến thức trọng tâm cho học sinh lớp 6, gồm: mục tiêu bài học, ghi nhớ ngắn gọn, ví dụ mẫu có lời giải, lỗi thường gặp và câu hỏi tự kiểm tra nhanh; dùng ngôn ngữ dễ hiểu, đúng chuẩn kiến thức môn Toán 6.",
        ],
    )
    add_label_block(doc, "Bước 3. Thiết kế Phần B - Luyện tập có phản hồi: ", "Phần B được xây dựng để học sinh luyện tập ngay sau khi học kiến thức. Mỗi bài tập cần có yêu cầu rõ ràng, vùng chọn hoặc ô trả lời, nút kiểm tra, phản hồi đúng/sai, gợi ý khi học sinh làm chưa đúng và đáp án tham khảo. Nhờ đó học sinh có thể tự sửa lỗi thay vì chỉ biết kết quả cuối cùng.")
    add_code_block(
        doc,
        [
            "Câu lệnh/prompt đã thực hiện:",
            "Viết Phần B - Luyện tập gồm các bài tập tương tác, có ô trả lời hoặc lựa chọn đáp án, nút kiểm tra, phản hồi đúng sai, gợi ý khi học sinh làm sai và đáp án gợi ý; sắp xếp câu hỏi từ nhận biết đến vận dụng.",
        ],
    )
    add_label_block(doc, "Bước 4. Thiết kế Phần C - Kiểm tra đánh giá: ", "Phần C dùng để học sinh tự kiểm tra mức độ hiểu bài sau khi đã học và luyện tập. Nội dung gồm thời gian làm bài, thang điểm, câu hỏi trắc nghiệm, tự luận ngắn, vận dụng, nút nộp bài, đáp án và nhận xét kết quả. Phần này giúp học sinh biết mình đạt mức nào và cần ôn lại nội dung nào.")
    add_code_block(
        doc,
        [
            "Câu lệnh/prompt đã thực hiện:",
            "Viết Phần C - Kiểm tra đánh giá gồm thời gian làm bài, thang điểm 10, câu hỏi trắc nghiệm, câu hỏi tự luận ngắn, câu hỏi vận dụng, đáp án, nút nộp bài và nhận xét kết quả theo mức độ hoàn thành của học sinh.",
        ],
    )
    add_label_block(doc, "Bước 5. Chuẩn hóa phần Kết quả và chức năng in phiếu: ", "Sau khi học sinh hoàn thành bài, website cần hiển thị điểm hoặc mức độ hoàn thành, nhận xét, gợi ý ôn tập và nút in/lưu kết quả. Đây là phần hỗ trợ học sinh tự đánh giá quá trình học, đồng thời giúp giáo viên có căn cứ trao đổi lại với học sinh.")
    add_code_block(
        doc,
        [
            "Câu lệnh/prompt đã thực hiện:",
            "Tạo phần Kết quả hiển thị thông tin học sinh, điểm từng phần, tổng điểm, nhận xét ngắn gọn, nội dung cần ôn lại, nút làm lại và nút in kết quả học tập.",
        ],
    )
    add_label_block(doc, "Bước 6. Kiểm tra sự thống nhất của phiếu sau khi tạo: ", "Tôi mở phiếu trên trình duyệt và kiểm tra lần lượt các tab A, B, C, Kết quả. Nếu phần nào thiếu mục tiêu, thiếu phản hồi, sai đáp án, sai thang điểm hoặc giao diện khó đọc, tôi quay lại chỉnh prompt hoặc file nguồn rồi tạo lại phiếu.")
    add_code_block(
        doc,
        [
            "Câu lệnh/prompt đã thực hiện:",
            "Kiểm tra phiếu học tập HTML: đủ Phần A, Phần B, Phần C và Kết quả; các tab chuyển đúng; câu hỏi có phản hồi; đáp án đúng; thang điểm đúng; nút nộp bài và nút in hoạt động; giao diện dễ đọc trên máy tính và điện thoại.",
        ],
    )
    add_label_block(
        doc,
        "Kết quả sau khi thực hiện các bước: ",
        "Sau khi chuẩn hóa, mỗi phiếu học tập điện tử không còn là một bản câu hỏi rời rạc mà trở thành một quy trình học hoàn chỉnh. Học sinh bắt đầu bằng việc đọc mục tiêu và kiến thức ở Phần A, chuyển sang luyện tập có phản hồi ở Phần B, hoàn thành kiểm tra ở Phần C, sau đó xem kết quả để tự điều chỉnh. Giáo viên có thể dùng cùng một cấu trúc này cho nhiều bài học khác nhau, chỉ cần thay nội dung chuyên môn trong các file nguồn.",
    )
    add_label_block(doc, "Hình ảnh minh họa cách thực hiện: ", "Các hình dưới đây được chụp trực tiếp từ phiếu học tập điện tử Bài 1 - Tập hợp trên website, minh họa cấu trúc ba phần A/B/C đã được chuẩn hóa.")
    add_image(doc, "skkn_assets/phieu-bai-1-phan-a.png", "Minh họa Biện pháp 2. Phần A - Kiến thức trọng tâm trong phiếu học tập điện tử", width=5.8)
    add_image(doc, "skkn_assets/phieu-bai-1-phan-b.png", "Minh họa Biện pháp 2. Phần B - Luyện tập có phản hồi trong phiếu học tập điện tử", width=5.8)
    add_image(doc, "skkn_assets/phieu-bai-1-phan-c.png", "Minh họa Biện pháp 2. Phần C - Kiểm tra đánh giá trong phiếu học tập điện tử", width=5.8)
    add_label_block(doc, "Sản phẩm đạt được: ", "Biện pháp 2 tạo ra mẫu cấu trúc chung cho toàn bộ hệ thống phiếu học tập điện tử Toán 6. Mỗi phiếu đều có môi trường học tập rõ ràng, các tab A/B/C/Kết quả thống nhất, câu hỏi có phản hồi, kiểm tra có thang điểm, kết quả có nhận xét. Nhờ đó học sinh dễ tự học theo trình tự, còn giáo viên dễ kiểm tra, chỉnh sửa và nhân rộng học liệu.")

    add_heading(doc, "3.3. Biện pháp 3: Thiết kế prompt mẫu hỗ trợ giáo viên tạo học liệu", 3)
    add_label_block(doc, "Mục đích: ", "Hỗ trợ giáo viên tạo phiếu học tập điện tử nhanh hơn, giảm thời gian định dạng, nhưng vẫn giữ quyền kiểm soát chuyên môn đối với kiến thức, câu hỏi, đáp án và thang điểm.")
    add_label_block(doc, "Môi trường thực hiện: ", "Biện pháp này được thực hiện trong môi trường ChatGPT và Codex. ChatGPT được dùng để xây dựng prompt mẫu, còn Codex được dùng để đọc thư mục bài học, tạo file HTML và kiểm tra sản phẩm trên website.")
    add_table(
        doc,
        ["Môi trường", "Cách sử dụng trong biện pháp 3", "Sản phẩm thu được"],
        [
            (
                "ChatGPT",
                "Trao đổi để hình thành prompt mẫu, quy định cấu trúc đầu vào, yêu cầu đầu ra và tiêu chí kiểm tra nội dung sư phạm.",
                "Có bộ câu lệnh rõ ràng để giáo viên sử dụng lại khi tạo phiếu học tập mới.",
            ),
            (
                "Codex",
                "Làm việc trong thư mục dự án, đọc các file lesson.md, PhanA, PhanB, PhanC; tạo file HTML; kiểm tra liên kết, tab, đáp án, thang điểm và giao diện.",
                "Có phiếu học tập điện tử hoàn chỉnh, thống nhất với website chung.",
            ),
            (
                "Website/Trình duyệt",
                "Mở file HTML đầu ra để kiểm tra sản phẩm sau khi dùng prompt mẫu; đối chiếu nội dung hiển thị với mục tiêu bài học.",
                "Bảo đảm prompt tạo ra sản phẩm dùng được trong dạy học thực tế.",
            ),
        ],
        [1900, 4700, 2500],
    )
    add_label_block(doc, "Cách thực hiện: ", "Tôi trình bày quy trình thiết kế và sử dụng prompt mẫu theo các bước sau. Mỗi bước có câu lệnh/prompt cụ thể để giáo viên có thể áp dụng lại khi tạo học liệu cho bài mới.")
    add_label_block(doc, "Bước 1. Xác định đầu vào của prompt: ", "Tôi quy định AI không tự tạo tùy ý mà phải đọc dữ liệu từ thư mục bài học. Mỗi thư mục gồm file thông tin chung lesson.md và ba file nội dung tương ứng với Phần A, Phần B, Phần C.")
    add_code_block(
        doc,
        [
            "Câu lệnh/prompt đã thực hiện:",
            "Thiết kế quy ước file nguồn cho mỗi bài học Toán 6 gồm lesson.md, BaiX-PhanA.md, BaiX-PhanB.md, BaiX-PhanC.md; AI chỉ tạo phiếu học tập dựa trên các file nguồn này.",
        ],
    )
    add_label_block(doc, "Bước 2. Viết prompt mẫu tổng quát: ", "Prompt tổng quát được dùng như câu lệnh chính khi giáo viên muốn chuyển nội dung bài học thành phiếu HTML. Câu lệnh phải ngắn, dễ nhớ, nhưng đủ rõ để Codex đọc đúng thư mục bài học.")
    add_code_block(
        doc,
        [
            "Câu lệnh/prompt đã thực hiện:",
            "Hoàn thành phiếu học tập [Tên thư mục bài học]. Đọc lesson.md, file Phần A, file Phần B, file Phần C; tạo file HTML tương tác đúng cấu trúc website, có đủ tab A/B/C/Kết quả, phản hồi, đáp án, thang điểm và nút in.",
        ],
    )
    add_label_block(doc, "Bước 3. Bổ sung yêu cầu kiểm soát chuyên môn trong prompt: ", "Để tránh sai lệch kiến thức, tôi thêm yêu cầu AI phải giữ đúng nội dung chuyên môn, không tự ý thay đáp án và không thay đổi thang điểm nếu giáo viên chưa yêu cầu.")
    add_code_block(
        doc,
        [
            "Câu lệnh/prompt đã thực hiện:",
            "Khi tạo phiếu, giữ nguyên chuẩn kiến thức, đáp án, thang điểm và yêu cầu đánh giá trong file nguồn; nếu phát hiện nội dung chưa rõ, hãy nêu lại phần cần giáo viên kiểm tra trước khi hoàn thiện.",
        ],
    )
    add_label_block(doc, "Bước 4. Dùng câu lệnh mẫu để tạo các phiếu cụ thể: ", "Sau khi thống nhất prompt, tôi áp dụng câu lệnh cho từng thư mục bài học trong Chương 1. Cách làm này giúp giáo viên chỉ cần thay tên thư mục, không phải viết lại yêu cầu từ đầu.")
    add_code_block(
        doc,
        [
            "Các câu lệnh đã thực hiện:",
            "Hoàn thành phiếu học tập Toan6-Chuong1-Bai1",
            "Hoàn thành phiếu học tập Toan6-Chuong1-Bai2",
            "Hoàn thành phiếu học tập Toan6-Chuong1-Bai3",
            "Hoàn thành phiếu học tập Toan6-Chuong1-Bai4",
            "Hoàn thành phiếu học tập Toan6-Chuong1-Bai5",
            "Hoàn thành phiếu học tập Toan6-Chuong1-Bai6",
            "Hoàn thành phiếu học tập Toan6-Chuong1-Bai7",
            "Hoàn thành phiếu học tập Toan6-Chuong1-OnTap",
        ],
    )
    add_label_block(doc, "Bước 5. Kiểm tra sản phẩm sau khi AI tạo: ", "Tôi mở file HTML trên trình duyệt để kiểm tra từng phần. Nếu sản phẩm chưa đạt yêu cầu, tôi chỉnh lại prompt hoặc chỉnh lại file nguồn rồi yêu cầu Codex tạo lại.")
    add_code_block(
        doc,
        [
            "Câu lệnh/prompt đã thực hiện:",
            "Rà soát phiếu HTML vừa tạo: kiểm tra trang tiêu đề, thông tin học sinh, Phần A, Phần B, Phần C, Kết quả, đáp án, thang điểm, nút chuyển tab, nút nộp bài, nút in và khả năng hiển thị trên trình duyệt.",
        ],
    )
    add_label_block(doc, "Hình ảnh minh họa cách thực hiện: ", "Các hình dưới đây minh họa kết quả của việc dùng prompt mẫu: từ danh sách bài học trên website TMath đến phiếu học tập HTML đầu ra.")
    add_image(doc, "skkn_assets/website-danh-sach-bai.png", "Minh họa Biện pháp 3. Danh sách bài học được tạo và quản lý thống nhất trên website TMath", width=5.8)
    add_image(doc, "skkn_assets/website-phieu-bai-1.png", "Minh họa Biện pháp 3. Phiếu học tập HTML là sản phẩm đầu ra trên website TMath sau khi dùng prompt mẫu", width=5.8)
    add_label_block(doc, "Sản phẩm đạt được: ", "Prompt mẫu trở thành quy trình hỗ trợ giáo viên tạo học liệu số một cách có kiểm soát. Giáo viên có thể nhân rộng cho bài học mới mà không phải thiết kế lại từ đầu.")

    add_heading(doc, "3.4. Biện pháp 4: Tổ chức hoạt động học sinh trên website", 3)
    add_label_block(doc, "Mục đích: ", "Phát triển năng lực tự học, giúp học sinh biết tự đặt mục tiêu, tự luyện tập, tự kiểm tra và duy trì thói quen học tập thường xuyên.")
    add_label_block(doc, "Môi trường thực hiện: ", "Biện pháp này được thực hiện trong môi trường website phiếu học tập đã xây dựng, có sự hỗ trợ của ChatGPT trong thiết kế hướng dẫn học tập và Codex trong hoàn thiện các chức năng tương tác.")
    add_table(
        doc,
        ["Môi trường", "Cách sử dụng trong biện pháp 4", "Sản phẩm thu được"],
        [
            (
                "ChatGPT",
                "Hỗ trợ viết hướng dẫn học sinh tự học theo trình tự: vào website, chọn bài, học Phần A, làm Phần B, kiểm tra Phần C, xem kết quả.",
                "Có quy trình học tập rõ ràng, phù hợp để giáo viên hướng dẫn trên lớp và giao về nhà.",
            ),
            (
                "Codex",
                "Hoàn thiện chức năng tương tác trong phiếu: nhập thông tin học sinh, chuyển tab, kiểm tra bài luyện tập, nộp bài, hiển thị kết quả và in phiếu.",
                "Học sinh có thể tự thao tác trên website mà không cần giáo viên giải thích lại nhiều lần.",
            ),
            (
                "Website/Trình duyệt",
                "Học sinh sử dụng trực tiếp trên máy tính hoặc thiết bị cá nhân để học, luyện tập, tự kiểm tra và xem nhận xét.",
                "Hình thành quy trình tự học có phản hồi ngay trên sản phẩm số.",
            ),
        ],
        [1900, 4700, 2500],
    )
    add_label_block(doc, "Cách thực hiện: ", "Tôi tổ chức hoạt động học sinh trên website theo các bước tuần tự dưới đây. Các bước này có thể dùng trong giờ học trên lớp hoặc khi giao nhiệm vụ tự học ở nhà.")
    add_label_block(doc, "Bước 1. Giáo viên giao nhiệm vụ học tập trên website: ", "Giáo viên thông báo bài học cần hoàn thành, đường link hoặc mã QR truy cập website, yêu cầu học sinh ghi rõ họ tên, lớp và mã bài trước khi làm.")
    add_code_block(
        doc,
        [
            "Câu lệnh/hướng dẫn đã thực hiện:",
            "Các em truy cập website Phiếu học tập Toán 6, nhập họ tên và lớp, chọn bài học được giao, học và làm bài theo đúng thứ tự các phần A, B, C.",
        ],
    )
    add_label_block(doc, "Bước 2. Học sinh học Phần A - Kiến thức trọng tâm: ", "Học sinh đọc mục tiêu, ghi nhớ, ví dụ mẫu và lỗi thường gặp. Giáo viên yêu cầu học sinh ghi lại ít nhất một nội dung trọng tâm hoặc một lỗi cần tránh.")
    add_code_block(
        doc,
        [
            "Câu lệnh/hướng dẫn đã thực hiện:",
            "Em đọc Phần A - Kiến thức trọng tâm, ghi lại kiến thức cần nhớ, xem ví dụ mẫu và đánh dấu nội dung còn chưa hiểu để hỏi giáo viên.",
        ],
    )
    add_label_block(doc, "Bước 3. Học sinh làm Phần B - Luyện tập: ", "Học sinh làm bài luyện tập trực tiếp trên website, bấm kiểm tra để nhận phản hồi. Với câu làm sai, học sinh đọc gợi ý, sửa lại và thử thêm lần nữa.")
    add_code_block(
        doc,
        [
            "Câu lệnh/hướng dẫn đã thực hiện:",
            "Em hoàn thành Phần B - Luyện tập; sau mỗi câu bấm Kiểm tra, đọc phản hồi đúng/sai, sửa lỗi và ghi lại dạng bài còn nhầm.",
        ],
    )
    add_label_block(doc, "Bước 4. Học sinh hoàn thành Phần C - Kiểm tra đánh giá: ", "Sau khi luyện tập, học sinh làm bài kiểm tra cuối bài, nộp bài và xem kết quả. Kết quả giúp học sinh tự biết mức độ hiểu bài của mình.")
    add_code_block(
        doc,
        [
            "Câu lệnh/hướng dẫn đã thực hiện:",
            "Em chuyển sang Phần C - Kiểm tra đánh giá, làm bài trong thời gian quy định, bấm Nộp bài và xem kết quả, sau đó đọc nhận xét của hệ thống.",
        ],
    )
    add_label_block(doc, "Bước 5. Giáo viên tổ chức phản hồi sau khi học sinh hoàn thành: ", "Giáo viên yêu cầu học sinh báo cáo kết quả, nêu câu sai hoặc nội dung chưa hiểu. Từ đó, giáo viên chọn một số lỗi phổ biến để chữa chung hoặc giao thêm bài phù hợp.")
    add_code_block(
        doc,
        [
            "Câu lệnh/hướng dẫn đã thực hiện:",
            "Sau khi có kết quả, em ghi lại tổng điểm, câu sai và nội dung cần hỏi lại. Giáo viên thu nhận phản hồi để chữa bài hoặc giao nhiệm vụ bổ sung.",
        ],
    )
    add_bullets(
        doc,
        [
            "Trong giờ học trên lớp, giáo viên dùng website để khởi động, củng cố kiến thức hoặc tổ chức luyện tập nhanh.",
            "Ở nhà, học sinh truy cập lại bài học để ôn tập, làm thêm bài và tự kiểm tra trước khi đến lớp.",
            "Đối với học sinh còn hạn chế, giáo viên giao bài theo từng phần nhỏ, ưu tiên kiến thức trọng tâm và bài luyện tập cơ bản.",
            "Đối với học sinh khá giỏi, giáo viên bổ sung câu hỏi vận dụng, yêu cầu giải thích cách làm hoặc tự tạo ví dụ tương tự.",
        ],
    )
    add_label_block(doc, "Hình ảnh minh họa cách thực hiện: ", "Hình dưới đây minh họa phần Kết quả sau khi học sinh nộp bài trên website, thể hiện thông tin học sinh, tổng điểm, nhận xét và bảng chấm.")
    add_image(doc, "skkn_assets/phieu-bai-1-ket-qua.png", "Minh họa Biện pháp 4. Kết quả sau khi học sinh nộp bài kiểm tra trên website", width=5.8)
    add_label_block(doc, "Sản phẩm đạt được: ", "Học sinh có công cụ học tập rõ ràng, dễ theo dõi, có thể học lại nhiều lần và từng bước hình thành thói quen tự học môn Toán.")

    add_heading(doc, "3.5. Biện pháp 5: Theo dõi tiến độ, phản hồi và điều chỉnh học liệu", 3)
    add_label_block(doc, "Mục đích: ", "Giúp giáo viên nắm được tình hình học tập của học sinh, phát hiện nội dung còn khó và điều chỉnh phiếu học tập cho phù hợp với thực tế lớp học.")
    add_label_block(doc, "Môi trường thực hiện: ", "Biện pháp này được thực hiện trong môi trường website, trình duyệt và Codex. Website hiển thị tiến độ học tập, còn Codex được dùng để chỉnh sửa học liệu khi giáo viên phát hiện nội dung cần cải tiến.")
    add_table(
        doc,
        ["Môi trường", "Cách sử dụng trong biện pháp 5", "Sản phẩm thu được"],
        [
            (
                "Website",
                "Hiển thị bài đã mở, bài gần nhất, phần trăm tiến độ, chuỗi ngày học và huy hiệu nỗ lực để học sinh tự theo dõi.",
                "Có dữ liệu học tập ban đầu giúp học sinh nhìn thấy tiến trình của mình.",
            ),
            (
                "ChatGPT",
                "Hỗ trợ phân tích phản hồi của học sinh và gợi ý cách viết lại câu hỏi, gợi ý, nhận xét hoặc hướng dẫn học tập.",
                "Có phương án điều chỉnh học liệu phù hợp hơn với khó khăn của học sinh.",
            ),
            (
                "Codex",
                "Chỉnh sửa trực tiếp các file nội dung, file HTML, CSS, JavaScript và cập nhật lại tài liệu SKKN khi cần bổ sung minh họa.",
                "Học liệu được cải tiến liên tục dựa trên phản hồi thực tế.",
            ),
        ],
        [1900, 4700, 2500],
    )
    add_label_block(doc, "Cách thực hiện: ", "Tôi thực hiện theo dõi, phản hồi và điều chỉnh học liệu theo các bước dưới đây. Quy trình này giúp website không chỉ là nơi giao bài mà còn là công cụ hỗ trợ giáo viên cải tiến dạy học.")
    add_label_block(doc, "Bước 1. Kích hoạt và kiểm tra phần theo dõi tiến độ trên website: ", "Website ghi nhận bài học sinh đã mở, bài học gần nhất, phần trăm hoàn thành chương, chuỗi ngày học và huy hiệu nỗ lực. Giáo viên hướng dẫn học sinh quan sát phần này để tự đặt mục tiêu học tiếp.")
    add_code_block(
        doc,
        [
            "Câu lệnh/prompt đã thực hiện:",
            "Tạo phần theo dõi tiến độ học tập trên trang chủ, gồm số bài đã mở, phần trăm hoàn thành, bài gần nhất, chuỗi ngày học, huy hiệu nỗ lực và nút tiếp tục học.",
        ],
    )
    add_label_block(doc, "Bước 2. Thu nhận phản hồi từ học sinh sau khi làm phiếu: ", "Sau khi học sinh làm bài, giáo viên ghi lại câu học sinh sai nhiều, phần học sinh thấy khó, nội dung học sinh chưa hiểu và nhận xét của học sinh về giao diện hoặc thao tác.")
    add_code_block(
        doc,
        [
            "Câu lệnh/hướng dẫn đã thực hiện:",
            "Sau khi hoàn thành phiếu, em ghi lại câu sai, phần chưa hiểu, nội dung cần giáo viên giải thích thêm và gửi lại phản hồi cho giáo viên.",
        ],
    )
    doc.add_page_break()
    add_label_block(doc, "Bước 3. Phân tích lỗi và xác định nội dung cần điều chỉnh: ", "Giáo viên đối chiếu phản hồi với kết quả trên lớp để xác định nguyên nhân: học sinh thiếu kiến thức nền, câu hỏi diễn đạt chưa rõ, ví dụ chưa đủ hoặc gợi ý chưa phù hợp.")
    add_code_block(
        doc,
        [
            "Câu lệnh/prompt đã thực hiện:",
            "Phân tích các lỗi học sinh thường gặp trong phiếu học tập, xác định nguyên nhân và đề xuất điều chỉnh phần kiến thức, bài luyện tập, gợi ý, đáp án hoặc nhận xét.",
        ],
    )
    add_label_block(doc, "Bước 4. Điều chỉnh học liệu bằng ChatGPT và Codex: ", "Từ kết quả phân tích, tôi dùng ChatGPT để viết lại gợi ý hoặc câu hỏi, sau đó dùng Codex để cập nhật vào file nguồn và file HTML tương ứng.")
    add_code_block(
        doc,
        [
            "Câu lệnh/prompt đã thực hiện:",
            "Điều chỉnh phiếu học tập Toan6-Chuong1-Bai1: bổ sung gợi ý cho câu học sinh sai nhiều, làm rõ yêu cầu bài tập, giữ nguyên đáp án và thang điểm, cập nhật lại file HTML.",
        ],
    )
    add_label_block(doc, "Bước 5. Kiểm tra lại sau khi điều chỉnh: ", "Sau khi sửa, tôi mở website trên trình duyệt để kiểm tra nội dung mới, thử các câu hỏi, kiểm tra phản hồi và bảo đảm học sinh vẫn học theo đúng trình tự.")
    add_code_block(
        doc,
        [
            "Câu lệnh/prompt đã thực hiện:",
            "Kiểm tra lại website sau khi điều chỉnh: tiến độ hiển thị đúng, bài học mở đúng, phản hồi đúng, đáp án đúng, nhận xét rõ ràng, giao diện không bị lỗi.",
        ],
    )
    add_bullets(
        doc,
        [
            "Bổ sung gợi ý ở những câu học sinh thường sai.",
            "Rút gọn phần lý thuyết nếu học sinh thấy quá dài, tăng ví dụ mẫu nếu học sinh chưa hiểu.",
            "Điều chỉnh câu hỏi kiểm tra theo mức độ nhận biết, thông hiểu, vận dụng.",
            "Cập nhật prompt mẫu khi phát hiện yêu cầu đầu ra chưa đủ rõ hoặc chưa phù hợp với lớp học.",
        ],
    )
    add_label_block(doc, "Hình ảnh minh họa cách thực hiện: ", "Hình dưới đây minh họa phần theo dõi tiến độ học tập trên website, gồm số bài đã mở, phần trăm hoàn thành, chuỗi ngày học, huy hiệu và bài gần nhất.")
    add_image(doc, "skkn_assets/website-tien-do.png", "Minh họa Biện pháp 5. Theo dõi tiến độ, chuỗi ngày học và huy hiệu nỗ lực trên website", width=5.8)
    add_label_block(doc, "Sản phẩm đạt được: ", "Học liệu không cố định một lần mà được cải tiến theo phản hồi thực tế, bảo đảm phù hợp hơn với năng lực và nhu cầu của học sinh.")


def add_application_and_meaning(doc):
    add_heading(doc, "4. Khả năng ứng dụng và triển khai của sáng kiến", 2)
    add_heading(doc, "4.1. Khả năng ứng dụng", 3)
    for text in [
        "Sáng kiến có thể áp dụng trong dạy học môn Toán lớp 6 ở các tiết học bài mới, luyện tập, ôn tập chương, phụ đạo học sinh còn hạn chế và bồi dưỡng học sinh khá giỏi. Website cũng có thể được dùng để giao nhiệm vụ học tập ở nhà hoặc chia sẻ cho phụ huynh theo dõi việc học của học sinh.",
        "Mô hình website tĩnh phù hợp với điều kiện cơ sở vật chất phổ biến vì không cần hệ thống máy chủ phức tạp. Giáo viên có thể chia sẻ đường link, mã QR hoặc mở trực tiếp file HTML trên máy tính, máy chiếu, tivi thông minh.",
        "Quy trình prompt mẫu có thể nhân rộng cho các chương khác của Toán 6 và các khối lớp khác. Khi đã có cấu trúc thư mục, giáo viên chỉ cần thay nội dung chuyên môn, giữ nguyên quy trình tạo phiếu và kiểm tra sản phẩm.",
    ]:
        add_body_paragraph(doc, text)
    add_heading(doc, "4.2. Điều kiện triển khai", 3)
    add_bullets(
        doc,
        [
            "Giáo viên cần chuẩn bị nội dung chuyên môn chính xác, gồm kiến thức trọng tâm, bài luyện tập, kiểm tra đánh giá, đáp án và thang điểm.",
            "Học sinh cần được hướng dẫn thao tác cơ bản: mở website, chọn bài, chuyển phần học, làm bài, xem phản hồi và in kết quả khi cần.",
            "Nhà trường cần hỗ trợ thiết bị trình chiếu, đường truyền internet và khuyến khích giáo viên chia sẻ học liệu số trong tổ chuyên môn.",
            "Giáo viên cần kiểm tra sản phẩm sau khi AI hỗ trợ tạo phiếu để bảo đảm không có sai lệch kiến thức, đáp án hoặc yêu cầu đánh giá.",
        ],
    )

    add_heading(doc, "5. Hướng phát triển của sáng kiến", 2)
    add_bullets(
        doc,
        [
            "Mở rộng website từ Chương 1 sang các chương còn lại của Toán 6, sau đó phát triển cho lớp 7, lớp 8 và lớp 9.",
            "Bổ sung ngân hàng câu hỏi theo mức độ nhận biết, thông hiểu, vận dụng và vận dụng cao.",
            "Hoàn thiện chức năng lưu kết quả kiểm tra để giáo viên có thể tổng hợp thuận tiện hơn khi tổ chức học tập trên lớp.",
            "Bổ sung hướng dẫn dành cho phụ huynh để hỗ trợ học sinh tự học ở nhà.",
            "Tiếp tục cải tiến prompt mẫu để giáo viên ở các môn học khác có thể sử dụng khi tạo phiếu học tập điện tử.",
        ],
    )

    add_heading(doc, "6. Ý nghĩa của sáng kiến", 2)
    add_heading(doc, "6.1. Đối với học sinh", 3)
    add_body_paragraph(
        doc,
        "Sáng kiến giúp học sinh có công cụ học tập rõ ràng, sinh động và dễ sử dụng. Các em biết học theo trình tự, được luyện tập, được tự kiểm tra và nhận phản hồi. Qua đó, học sinh từng bước hình thành thói quen tự học, tự đánh giá và tự điều chỉnh lỗi sai."
    )
    add_heading(doc, "6.2. Đối với giáo viên", 3)
    add_body_paragraph(
        doc,
        "Sáng kiến hỗ trợ giáo viên đổi mới phương pháp dạy học, giảm thời gian thiết kế phiếu học tập thủ công, chuẩn hóa quy trình tạo học liệu và nâng cao năng lực ứng dụng AI trong giảng dạy. Giáo viên có thêm công cụ để tổ chức hoạt động học tập linh hoạt, kiểm soát kiến thức và hỗ trợ học sinh kịp thời."
    )
    add_heading(doc, "6.3. Đối với nhà trường", 3)
    add_body_paragraph(
        doc,
        "Sáng kiến góp phần xây dựng môi trường học tập hiện đại, thúc đẩy chuyển đổi số trong dạy học và tạo nguồn học liệu số có thể chia sẻ trong tổ chuyên môn. Sản phẩm có khả năng nhân rộng, dễ cập nhật và phù hợp với định hướng đổi mới giáo dục hiện nay."
    )


def add_conclusion(doc):
    add_section_heading(doc, "III. KẾT LUẬN")
    add_heading(doc, "1. Bài học kinh nghiệm", 2)
    add_bullets(
        doc,
        [
            "Cần bắt đầu từ nhu cầu học tập thực tế của học sinh, không chạy theo hình thức công nghệ. Website phải giúp học sinh học dễ hơn, luyện tập nhiều hơn và tự kiểm tra được kết quả.",
            "Nội dung chuyên môn là yếu tố quyết định chất lượng phiếu học tập. AI chỉ là công cụ hỗ trợ chuyển hóa và trình bày nội dung, giáo viên vẫn phải kiểm tra kiến thức, đáp án và thang điểm.",
            "Prompt mẫu cần cụ thể, có quy trình rõ ràng và yêu cầu đầu ra chặt chẽ. Khi prompt rõ, sản phẩm tạo ra ổn định hơn và giáo viên dễ kiểm soát hơn.",
            "Cần hướng dẫn học sinh sử dụng website từng bước, đặc biệt với những em còn hạn chế về kỹ năng công nghệ hoặc chưa có thói quen tự học.",
            "Học liệu số cần được điều chỉnh thường xuyên dựa trên phản hồi của học sinh, kết quả luyện tập và trao đổi trong tổ chuyên môn.",
        ],
    )
    add_heading(doc, "2. Những kiến nghị, đề xuất", 2)
    add_heading(doc, "2.1. Đối với nhà trường", 3)
    add_bullets(
        doc,
        [
            "Tạo điều kiện về thiết bị, đường truyền internet và không gian sinh hoạt chuyên môn để giáo viên xây dựng, chia sẻ học liệu số.",
            "Khuyến khích giáo viên ứng dụng AI có kiểm soát trong thiết kế bài học, phiếu học tập và kiểm tra đánh giá.",
            "Tổ chức các buổi trao đổi kinh nghiệm về xây dựng website học tập, sử dụng prompt mẫu và kiểm duyệt nội dung số.",
        ],
    )
    add_heading(doc, "2.2. Đối với tổ chuyên môn", 3)
    add_bullets(
        doc,
        [
            "Cùng góp ý nội dung phiếu học tập, đặc biệt là câu hỏi luyện tập, kiểm tra đánh giá, đáp án và thang điểm.",
            "Xây dựng thư viện prompt mẫu, thư viện bài học và ngân hàng câu hỏi dùng chung cho môn Toán.",
            "Thống nhất tiêu chí kiểm tra chất lượng học liệu số trước khi đưa vào sử dụng với học sinh.",
        ],
    )
    add_heading(doc, "2.3. Đối với giáo viên", 3)
    add_bullets(
        doc,
        [
            "Chủ động học hỏi công nghệ, nhưng luôn đặt mục tiêu sư phạm và sự phù hợp với học sinh lên trước.",
            "Sử dụng AI như công cụ hỗ trợ, không sao chép máy móc kết quả AI khi chưa kiểm tra.",
            "Thường xuyên lấy ý kiến học sinh để cải tiến nội dung, giao diện và mức độ bài tập.",
        ],
    )
    add_body_paragraph(
        doc,
        "Sáng kiến “Xây dựng website phiếu học tập điện tử môn Toán 6 nhằm phát triển năng lực tự học cho học sinh” có ý nghĩa thiết thực trong đổi mới dạy học. Sản phẩm giúp học sinh học tập chủ động hơn, giúp giáo viên thiết kế học liệu số khoa học hơn và góp phần thực hiện chuyển đổi số trong nhà trường.",
    )
    doc.add_paragraph()
    add_centered_paragraph(doc, "Ngày 15 tháng 01 năm 2027", size=13, after=4)
    add_centered_paragraph(doc, "Người viết sáng kiến", size=13, bold=True, after=46)
    add_centered_paragraph(doc, "Nguyễn Thị Huyền Trang", size=13, bold=True, after=0)


def add_appendix(doc):
    doc.add_page_break()
    add_section_heading(doc, "PHỤ LỤC")
    add_heading(doc, "PL1. Cấu trúc sản phẩm website phiếu học tập điện tử", 2)
    add_body_paragraph(
        doc,
        "Sản phẩm website được tổ chức theo cấu trúc thư mục rõ ràng để giáo viên dễ quản lý, dễ cập nhật và dễ nhân rộng cho các bài học mới."
    )
    rows = [
        ("index.html", "Trang chủ website, danh sách bài học, hướng dẫn, đăng nhập và tiến độ học tập."),
        ("assets/site.css, assets/site.js", "Giao diện và chức năng dùng chung cho trang chủ."),
        ("assets/lesson-shell.css, assets/lesson-shell.js", "Thanh điều hướng chung cho từng bài học: trang chủ, bài trước, bài sau, in phiếu."),
        ("templates/base_prompt.md", "Prompt nền hướng dẫn AI tạo phiếu học tập tương tác từ nội dung giáo viên chuẩn bị."),
        ("promt.md", "Lệnh nền và quy trình xử lý khi giáo viên yêu cầu hoàn thành phiếu học tập."),
        ("Toan6-Chuong1-BaiX/", "Thư mục từng bài gồm lesson.md, PhanA, PhanB, PhanC và file HTML đầu ra."),
    ]
    add_table(doc, ["Thành phần", "Vai trò"], rows, [3300, 5800])

    add_heading(doc, "PL2. Prompt mẫu hỗ trợ giáo viên", 2)
    add_body_paragraph(
        doc,
        "Giáo viên có thể sử dụng prompt mẫu sau để yêu cầu AI tạo hoặc cập nhật phiếu học tập điện tử. Nội dung chuyên môn cần được giáo viên chuẩn bị và kiểm tra trước khi sử dụng với học sinh."
    )
    add_code_block(
        doc,
        [
            "Bạn là agent tạo phiếu học tập tương tác từ nội dung Markdown của giáo viên.",
            "Khi nhận lệnh: Hoàn thành phiếu học tập [Tên thư mục]",
            "Hãy đọc lesson.md, Phần A - kiến thức trọng tâm, Phần B - luyện tập, Phần C - kiểm tra đánh giá.",
            "Tạo một file HTML độc lập trong thư mục bài học.",
            "HTML phải có đủ: thông tin học sinh, điều hướng, Phần A, Phần B, Phần C, kết quả và nút in.",
            "Giữ đúng kiến thức, câu hỏi, đáp án và thang điểm trong file nguồn.",
            "Nếu thiếu dữ liệu hoặc file còn là bản mẫu, báo rõ phần thiếu để giáo viên bổ sung.",
        ],
    )
    add_heading(doc, "PL3. Minh họa sản phẩm", 2)
    add_image(doc, "assets/home-hero.png", "Hình 1. Hình ảnh minh họa định hướng học tập Toán 6 trên website", width=6.0)
    add_image(doc, "assets/learning-cards-static.png", "Hình 2. Ba hoạt động chính của phiếu học tập: kiến thức trọng tâm, luyện tập, kiểm tra kiến thức", width=5.8)

    doc.add_page_break()
    add_heading(doc, "PL4. Minh họa hoạt động của phiếu học tập điện tử", 2)
    add_body_paragraph(
        doc,
        "Trang phụ lục này minh họa cách phiếu học tập điện tử hoạt động khi học sinh trực tiếp thao tác trên website. Học sinh chọn phần học, trả lời câu hỏi, bấm nút kiểm tra và nhận phản hồi ngay trên phiếu. Đây là điểm khác biệt của phiếu học tập điện tử so với phiếu giấy thông thường."
    )
    add_bullets(
        doc,
        [
            "Học sinh chuyển giữa các phần A - Kiến thức, B - Luyện tập, C - Kiểm tra và Kết quả.",
            "Ở phần luyện tập, học sinh chọn đáp án hoặc nhập câu trả lời rồi bấm nút kiểm tra.",
            "Hệ thống phản hồi đúng/sai, đưa ra gợi ý hoặc đáp án để học sinh tự sửa lỗi.",
            "Sau phần kiểm tra, học sinh nộp bài, xem kết quả, nhận xét và có thể in phiếu khi cần.",
        ],
    )
    add_image(doc, "skkn_assets/phieu-bai-1-hoat-dong-luyen-tap.png", "Hình 3. Hoạt động luyện tập trên phiếu: học sinh chọn đáp án, bấm kiểm tra và nhận phản hồi", width=5.8)

    doc.add_page_break()
    add_heading(doc, "PL5. Minh họa các khu vực chính trên website TMath", 2)
    add_body_paragraph(
        doc,
        "Trang phụ lục này bổ sung hình ảnh minh họa các khu vực học sinh thường sử dụng trên website TMath, gồm ô đăng nhập, hướng dẫn học tập, kiến thức trọng tâm, luyện tập và kiểm tra. Các hình được chụp trực tiếp từ sản phẩm website sau khi đổi tên thành TMath."
    )
    add_image(doc, "skkn_assets/website-dang-nhap.png", "Hình 4. Ô đăng nhập và đăng kí tài khoản học sinh trên website TMath", width=5.8)
    add_image(doc, "skkn_assets/website-huong-dan.png", "Hình 5. Khu vực hướng dẫn học tập trên website TMath", width=5.8)

    doc.add_page_break()
    add_image(doc, "skkn_assets/phieu-bai-1-phan-a.png", "Hình 6. Phần A - Kiến thức trọng tâm trong phiếu học tập điện tử", width=5.8)
    add_image(doc, "skkn_assets/phieu-bai-1-phan-b.png", "Hình 7. Phần B - Luyện tập trong phiếu học tập điện tử", width=5.8)

    doc.add_page_break()
    add_image(doc, "skkn_assets/phieu-bai-1-phan-c.png", "Hình 8. Phần C - Kiểm tra trong phiếu học tập điện tử", width=5.8)


def build():
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_intro(doc)
    add_theory_and_status(doc)
    add_measures(doc)
    add_application_and_meaning(doc)
    add_conclusion(doc)
    add_appendix(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
